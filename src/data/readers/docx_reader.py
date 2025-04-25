"""DOCX document reader."""

import os
from pathlib import Path
from typing import Dict, Any, List, Optional
import docx
from datetime import datetime

from .base_reader import BaseReader
from ...core.exceptions.exceptions import DocumentProcessingError
from ...core.logging.logger import get_logger

logger = get_logger(__name__)

class DocxReader(BaseReader):
    """Reader for DOCX documents."""
    
    def supports(self, file_path: str) -> bool:
        """Check if this reader supports the given file type."""
        return self._get_file_extension(file_path) == "docx"
    
    def read(self, file_path: str) -> Dict[str, Any]:
        """Read a DOCX document and return its content and metadata."""
        try:
            logger.info(f"Reading DOCX document: {file_path}")
            
            # Check if file exists
            if not os.path.exists(file_path):
                raise DocumentProcessingError(f"File not found: {file_path}")
            
            # Open the document
            doc = docx.Document(file_path)
            
            # Extract text
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():  # Skip empty cells
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            # Join paragraphs with newlines
            text = "\n\n".join(paragraphs)
            
            # Extract metadata
            metadata = self._extract_metadata(doc, file_path)
            
            logger.info(f"Successfully read DOCX document: {file_path}")
            return {
                "text": text,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error reading DOCX document {file_path}: {str(e)}")
            raise DocumentProcessingError(f"Error reading DOCX document: {str(e)}")
    
    def _extract_metadata(self, doc: docx.Document, file_path: str) -> Dict[str, Any]:
        """Extract metadata from a DOCX document."""
        metadata = {}
        
        # File metadata
        file_path_obj = Path(file_path)
        metadata["filename"] = file_path_obj.name
        metadata["file_extension"] = "docx"
        metadata["file_path"] = str(file_path_obj.absolute())
        metadata["file_size"] = file_path_obj.stat().st_size
        metadata["last_modified"] = datetime.fromtimestamp(
            file_path_obj.stat().st_mtime
        ).isoformat()
        
        # Document properties
        try:
            core_properties = doc.core_properties
            metadata["title"] = core_properties.title if core_properties.title else ""
            metadata["author"] = core_properties.author if core_properties.author else ""
            metadata["created"] = core_properties.created.isoformat() if core_properties.created else ""
            metadata["modified"] = core_properties.modified.isoformat() if core_properties.modified else ""
            metadata["last_modified_by"] = core_properties.last_modified_by if core_properties.last_modified_by else ""
            metadata["revision"] = core_properties.revision if core_properties.revision else ""
            metadata["category"] = core_properties.category if core_properties.category else ""
            metadata["comments"] = core_properties.comments if core_properties.comments else ""
            metadata["subject"] = core_properties.subject if core_properties.subject else ""
            metadata["keywords"] = core_properties.keywords if core_properties.keywords else ""
        except Exception as e:
            logger.warning(f"Error extracting core properties from {file_path}: {str(e)}")
        
        # Document statistics
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)
        metadata["section_count"] = len(doc.sections)
        
        # Try to extract document type from content or filename
        metadata["document_type"] = self._guess_document_type(doc, file_path_obj.name)
        
        return metadata
    
    def _guess_document_type(self, doc: docx.Document, filename: str) -> str:
        """Guess the document type based on content and filename."""
        # Check first few paragraphs for type indicators
        first_paragraphs = " ".join([p.text for p in doc.paragraphs[:10] if p.text.strip()])
        
        # Check for regulatory document indicators
        if any(term in first_paragraphs.lower() for term in ["regulation", "directive", "advisory", "circular", "order"]):
            return "regulatory"
        
        # Check for accident report indicators
        if any(term in first_paragraphs.lower() for term in ["accident", "incident", "investigation", "safety report"]):
            return "accident_report"
        
        # Check for manual indicators
        if any(term in first_paragraphs.lower() for term in ["manual", "handbook", "guide", "procedure"]):
            return "manual"
        
        # Check filename for indicators
        filename_lower = filename.lower()
        if any(term in filename_lower for term in ["faa", "easa", "regulation", "part", "advisory"]):
            return "regulatory"
        if any(term in filename_lower for term in ["accident", "incident", "investigation", "safety"]):
            return "accident_report"
        if any(term in filename_lower for term in ["manual", "handbook", "guide", "procedure"]):
            return "manual"
        
        # Default
        return "unknown"
